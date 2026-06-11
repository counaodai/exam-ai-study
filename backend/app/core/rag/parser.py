import os
import re
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass, field

import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass
class ParsedPage:
    page_number: int
    content: str


@dataclass
class ParsedDocument:
    filename: str
    file_type: str
    pages: list[ParsedPage] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        return "\n\n".join(page.content for page in self.pages if page.content.strip())


class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        ...


class PDFParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        filename = os.path.basename(file_path)
        
        # 优先尝试使用 pdfminer.six 解析（效果更好）
        try:
            return self._parse_with_pdfminer(filename, file_path)
        except ImportError:
            logger.info("pdfminer.six 未安装，使用 PyPDF2 解析 PDF")
        except Exception as e:
            logger.warning(f"pdfminer.six 解析失败: {e}，回退到 PyPDF2")
        
        # 回退到 PyPDF2
        return self._parse_with_pypdf2(filename, file_path)
    
    def _parse_with_pypdf2(self, filename: str, file_path: str) -> ParsedDocument:
        """使用 PyPDF2 解析 PDF。"""
        pages: list[ParsedPage] = []
        
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                text = self._clean_text(text)
                if text.strip():
                    pages.append(ParsedPage(page_number=i + 1, content=text))
        
        return ParsedDocument(filename=filename, file_type="pdf", pages=pages)
    
    def _parse_with_pdfminer(self, filename: str, file_path: str) -> ParsedDocument:
        """使用 pdfminer.six 解析 PDF（效果更好，保留段落结构）。"""
        from io import BytesIO
        from pdfminer.high_level import extract_text
        from pdfminer.layout import LAParams
        
        la_params = LAParams(
            line_margin=0.5,       # 降低行间距阈值，更好的段落识别
            word_margin=0.1,       # 降低词间距阈值
            char_margin=2.0,       # 字符间距阈值
            boxes_flow=0.5,        # 布局分析
        )
        
        text = extract_text(file_path, laparams=la_params)
        
        # 将全文按页分割（假设每 2000 字符为一页，实际更准确的方式是手动分割）
        pages: list[ParsedPage] = []
        if text.strip():
            cleaned_text = self._clean_text(text)
            pages.append(ParsedPage(page_number=1, content=cleaned_text))
        
        return ParsedDocument(filename=filename, file_type="pdf", pages=pages)
    
    def _clean_text(self, text: str) -> str:
        """清理 PDF 解析后的文本，保留完整结构。
        
        关键改进：
        1. 不移除空行，因为空行是 PDF 结构的重要分隔符
        2. 只清理不可见字符（如 \x00 零宽字符）
        3. 压缩连续空行但不完全删除
        """
        # 清理不可见字符
        text = text.replace("\x00", "")
        
        # 替换连续多个换行为双换行（保留段落分隔）
        text = re.sub(r'[\r\n]+', '\n', text)
        
        # 保留空行，但压缩连续 3 个以上换行为 2 个
        lines = text.split("\n")
        cleaned = []
        blank_count = 0
        for line in lines:
            stripped = line.strip()
            if not stripped:
                blank_count += 1
                if blank_count <= 2:  # 最多保留 2 个连续空行
                    cleaned.append("")
            else:
                blank_count = 0
                cleaned.append(stripped)
        
        result = "\n".join(cleaned)
        
        # 如果文本被严重截断（单行超过 5000 字符），尝试按字符位置重新分行
        if "\n" not in result and len(result) > 5000:
            result = self._split_long_line(result)
        
        return result
    
    def _split_long_line(self, text: str) -> str:
        """尝试将超长行按句子边界拆分。"""
        # 在句号、问号、感叹号后换行
        text = re.sub(r'([。！？])', r'\1\n', text)
        # 在换行符后换行
        text = re.sub(r'([\n\r])', r'\n', text)
        return text


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        filename = os.path.basename(file_path)
        doc = DocxDocument(file_path)
        paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        content = "\n".join(paragraphs)
        pages = [ParsedPage(page_number=1, content=content)] if content else []
        return ParsedDocument(filename=filename, file_type="docx", pages=pages)


class TextParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        filename = os.path.basename(file_path)
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        pages = [ParsedPage(page_number=1, content=content)] if content.strip() else []
        return ParsedDocument(filename=filename, file_type="txt", pages=pages)


class MarkdownParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        filename = os.path.basename(file_path)
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        pages = [ParsedPage(page_number=1, content=content)] if content.strip() else []
        return ParsedDocument(filename=filename, file_type="md", pages=pages)


PARSER_MAP: dict[str, BaseParser] = {
    "pdf": PDFParser(),
    "docx": DocxParser(),
    "doc": DocxParser(),
    "txt": TextParser(),
    "md": MarkdownParser(),
}


def get_parser(file_type: str) -> BaseParser:
    parser = PARSER_MAP.get(file_type.lower())
    if not parser:
        raise ValueError(f"不支持的文件类型: {file_type}")
    return parser


def parse_document(file_path: str, file_type: str) -> ParsedDocument:
    parser = get_parser(file_type)
    return parser.parse(file_path)
